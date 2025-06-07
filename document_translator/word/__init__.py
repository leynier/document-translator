from pathlib import Path
from typing import Annotated

from docx import Document
from typer import Option, Typer

from ..common.translators import Translator, get_translator
from .styler import copy_paragraph_style, copy_run_style

app = Typer(no_args_is_help=True)


@app.callback()
def main(
    input: Annotated[
        Path,
        Option(
            "--input",
            "-i",
            exists=True,
            dir_okay=False,
            resolve_path=True,
            help="Input file path",
        ),
    ],
    output: Annotated[
        str,
        Option(
            "--output",
            "-o",
            help="Output file path",
        ),
    ],
    language_source: Annotated[
        str,
        Option(
            "--language-source",
            "-s",
            help="Source language. For example: 'spanish'",
        ),
    ],
    language_target: Annotated[
        str,
        Option(
            "--language-target",
            "-t",
            help="Target language. For example: 'english'",
        ),
    ],
    translator: Annotated[
        Translator,
        Option(
            help="Translator",
            case_sensitive=False,
        ),
    ] = Translator.openai,
):
    translate = get_translator(translator)
    texts: list[str] = []
    doc = Document(str(input))
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs):
        print(f"Paragraph {i + 1} of {len(doc.paragraphs)}")
        if len(para.runs) == 1:
            text = para.runs[0].text
            if text.strip():
                texts.append(text)
            continue
        text = para.text
        if text.strip():
            texts.append(text)
    step = max(1, len(texts) // 100)
    print(f"Total texts: {len(texts)}")
    texts = [
        text
        for i in range(0, len(texts), step)
        for text in translate(
            texts[i : i + step],
            language_source,
            language_target,
        )
    ]
    print(f"Total translated texts: {len(texts)}")
    translated_doc = Document()
    index = 0
    for para in doc.paragraphs:
        if len(para.runs) == 1:
            text = para.runs[0].text
            if text.strip():
                translated_text = texts[index]
                index += 1
            else:
                translated_text = text
            new_para = translated_doc.add_paragraph()
            copy_paragraph_style(para, new_para)
            new_run = new_para.add_run(translated_text)
            copy_run_style(para.runs[0], new_run)
            continue
        text = para.text
        if text.strip():
            translated_text = texts[index]
            index += 1
        else:
            translated_text = text
        new_para = translated_doc.add_paragraph(translated_text)
        copy_paragraph_style(para, new_para)
    Path(output).parent.mkdir(parents=True, exist_ok=True)
    if not output.endswith(".docx"):
        output += ".docx"
    translated_doc.save(output)
