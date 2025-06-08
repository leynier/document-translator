# ruff: noqa: E402
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

os.environ.setdefault("OPENAI_API_KEY", "test")

from doculingo.word.styler import copy_paragraph_style, copy_run_style


def test_copy_run_style() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("hello")
    run.bold = True
    run.italic = True
    run.underline = True
    run.font.size = Pt(16)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x11, 0x22, 0x33)

    new_doc = Document()
    new_run = new_doc.add_paragraph().add_run("world")
    copy_run_style(run, new_run)

    assert new_run.bold is True
    assert new_run.italic is True
    assert new_run.underline is True
    assert new_run.font.size == Pt(16)
    assert new_run.font.name == "Arial"
    assert new_run.font.color.rgb == RGBColor(0x11, 0x22, 0x33)


def test_copy_paragraph_style() -> None:
    doc = Document()
    para = doc.add_paragraph("text")
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 2
    para.paragraph_format.keep_together = True
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.page_break_before = True
    para.paragraph_format.widow_control = True
    para.style = doc.styles["Title"]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc = Document()
    new_para = new_doc.add_paragraph()
    copy_paragraph_style(para, new_para)

    fmt1 = para.paragraph_format
    fmt2 = new_para.paragraph_format

    assert fmt2.left_indent == fmt1.left_indent
    assert fmt2.right_indent == fmt1.right_indent
    assert fmt2.space_before == fmt1.space_before
    assert fmt2.space_after == fmt1.space_after
    assert fmt2.line_spacing == fmt1.line_spacing
    assert fmt2.keep_together == fmt1.keep_together
    assert fmt2.keep_with_next == fmt1.keep_with_next
    assert fmt2.page_break_before == fmt1.page_break_before
    assert fmt2.widow_control == fmt1.widow_control
    assert new_para.style.name == para.style.name
    assert new_para.alignment == para.alignment
